"""Document CRUD and Aenderungsantrag CRUD endpoints"""
import asyncio
import os
import tempfile
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import APIRouter, Depends, HTTPException, status, Query, UploadFile, File, Request
from fastapi.responses import Response, FileResponse
from sqlalchemy.orm import Session, joinedload
from typing import List, Optional

import aiofiles

from app.api.deps import get_db
from app.core.rbac import require_role
from app.services.pdf import docx_to_pdf
from app.models.document import Document
from app.models.document_aenderungsantrag import DocumentAenderungsantrag
from app.models.document_aenderung import DocumentAenderung
from app.models.email_template import EmailTemplate
from app.models.user import User
from app.services.email import send_email, render_template
from app.schemas.document import (
    DocumentCreate,
    DocumentUpdate,
    DocumentResponse,
    DocumentAenderungsantragCreate,
    DocumentAenderungsantragUpdate,
    DocumentAenderungsantragResponse,
    DocumentAenderungCreate,
    DocumentAenderungUpdate,
    DocumentAenderungResponse,
)
from app.config import settings
from app.services.audit import log_action

router = APIRouter()

DOCUMENT_UPLOAD_DIR = os.path.join(settings.UPLOAD_DIR, "dokumente")

ALLOWED_DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc", ".odt", ".txt"}
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB

def _validate_upload(filename: str, content_length: int = 0) -> str:
    """Validate file extension and return safe extension."""
    ext = os.path.splitext(filename)[1].lower() if filename else ""
    if ext not in ALLOWED_DOCUMENT_EXTENSIONS:
        raise HTTPException(status_code=400, detail=f"Dateityp '{ext}' nicht erlaubt. Erlaubt: {', '.join(ALLOWED_DOCUMENT_EXTENSIONS)}")
    return ext

def _safe_file_path(base_dir: str, file_path: str) -> bool:
    """Check that file_path is within base_dir (prevent path traversal)."""
    real_base = os.path.realpath(base_dir)
    real_path = os.path.realpath(file_path)
    return real_path.startswith(real_base)


def _stellen_for_export(antrag: DocumentAenderungsantrag) -> List[dict]:
    """Ergibt eine Liste von Stellen (dict mit bezug, alte_fassung, neue_fassung, aenderungstext) für Export. Legacy: ein Eintrag aus antrag wenn keine stellen."""
    if antrag.stellen:
        return [
            {
                "bezug": s.bezug or "",
                "alte_fassung": s.alte_fassung or "",
                "neue_fassung": s.neue_fassung or "",
                "aenderungstext": s.aenderungstext or _default_aenderungstext(s.bezug, s.alte_fassung, s.neue_fassung),
            }
            for s in sorted(antrag.stellen, key=lambda x: (x.position, x.id))
        ]
    # Legacy: eine Stelle aus dem Antrag
    return [
        {
            "bezug": "",
            "alte_fassung": antrag.alte_fassung or "",
            "neue_fassung": antrag.neue_fassung or "",
            "aenderungstext": antrag.antrag_text or _default_aenderungstext("", antrag.alte_fassung, antrag.neue_fassung),
        }
    ]


def _default_aenderungstext(bezug: str, alte: Optional[str], neue: Optional[str]) -> str:
    """Erzeugt einen Standard-Änderungstext wenn keiner gesetzt ist."""
    ref = f"{bezug} " if bezug else ""
    if not alte and neue:
        return f"In {ref}wird eingefügt:\n{neue}" if ref else f"Einfügung:\n{neue}"
    if alte and not neue:
        return f"{ref}wird gestrichen." if ref else "Streichung."
    if alte and neue:
        return f"{ref}wird wie folgt geändert:\n{neue}" if ref else f"Ersetzung:\n{neue}"
    return ""


def _send_aenderungsantrag_emails(db: Session, antrag: DocumentAenderungsantrag, doc: Document) -> None:
    """Sendet Benachrichtigungs-E-Mails für einen Änderungsantrag an konfigurierte Empfänger."""
    if not settings.email_configured:
        return
    recipients = settings.document_amendment_notify_emails_list
    if not recipients:
        return
    stellen = _stellen_for_export(antrag)
    stellen_uebersicht = "\n\n".join(
        (st.get("aenderungstext") or "").strip() or f"Alte Fassung: {st.get('alte_fassung', '')}\nNeue Fassung: {st.get('neue_fassung', '')}"
        for st in stellen
    )
    template_vars = {
        "dokument_titel": doc.titel or "",
        "antragsteller": antrag.antragsteller or "",
        "antrag_text": antrag.antrag_text or "",
        "begruendung": antrag.begruendung or "",
        "link": f"{settings.APP_URL}/dokumente/satzung/{doc.id}",
        "stellen_uebersicht": stellen_uebersicht,
    }
    templates = (
        db.query(EmailTemplate)
        .filter(
            EmailTemplate.scenario == "aenderungsantrag",
            EmailTemplate.typ == "benachrichtigung",
            EmailTemplate.kreisverband_id.is_(None),
        )
        .all()
    )
    for template in templates:
        subject = render_template(template.betreff, template_vars)
        body = render_template(template.inhalt, template_vars)
        send_email(to=recipients, subject=subject, body=body)


def _build_aenderungsantrag_docx(antrag: DocumentAenderungsantrag, doc: Document, out_path: str) -> None:
    """Erstellt eine DOCX-Datei: Titel, Antragsteller, Begründung, Änderungstext, Synopse (Tabelle)."""
    d = DocxDocument()
    style = d.styles["Normal"]
    style.font.size = Pt(11)
    # Titel
    title = d.add_paragraph()
    title.add_run("Änderungsantrag").bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_paragraph(doc.titel)
    if antrag.titel:
        d.add_paragraph(antrag.titel)
    d.add_paragraph()
    d.add_paragraph(f"Antragsteller: {antrag.antragsteller}")
    if antrag.begruendung:
        d.add_paragraph()
        p = d.add_paragraph()
        p.add_run("Begründung:").bold = True
        d.add_paragraph(antrag.begruendung)
    stellen = _stellen_for_export(antrag)
    # Änderungstext
    d.add_paragraph()
    h = d.add_paragraph()
    h.add_run("Änderungstext").bold = True
    h.paragraph_format.space_before = Pt(12)
    for i, st in enumerate(stellen):
        if st["bezug"]:
            d.add_paragraph(st["bezug"]).paragraph_format.space_before = Pt(6)
        d.add_paragraph(st["aenderungstext"])
    # Synopse
    d.add_paragraph()
    h2 = d.add_paragraph()
    h2.add_run("Synopse").bold = True
    h2.paragraph_format.space_before = Pt(18)
    table = d.add_table(rows=1 + len(stellen), cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Bezug"
    hdr[1].text = "Alte Fassung"
    hdr[2].text = "Neue Fassung"
    for i, st in enumerate(stellen):
        row = table.rows[i + 1].cells
        row[0].text = st["bezug"] or "—"
        row[1].text = st["alte_fassung"] or "—"
        row[2].text = st["neue_fassung"] or "—"
    d.save(out_path)


# ---------------------------------------------------------------------------
# Document CRUD
# ---------------------------------------------------------------------------

@router.get("/", response_model=List[DocumentResponse])
async def list_documents(
    typ: Optional[str] = Query(None, description="Filter by type"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """List all documents. Vorstand+ can view."""
    query = db.query(Document)
    if typ:
        query = query.filter(Document.typ == typ)
    query = query.order_by(Document.titel)
    return query.all()


@router.post("/", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def create_document(
    data: DocumentCreate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """Create a new document. Leitung+ can create."""
    doc = Document(
        titel=data.titel,
        typ=data.typ,
        aktueller_text=data.aktueller_text,
        version=data.version,
        gueltig_ab=data.gueltig_ab,
    )
    db.add(doc)
    db.flush()
    log_action(db, current_user.id, "create", "document", doc.id, f"Dokument erstellt: {doc.titel}", request)
    db.commit()
    db.refresh(doc)
    return doc


@router.get("/{document_id}", response_model=DocumentResponse)
async def get_document(
    document_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Get a single document by ID. Vorstand+ can view."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    return doc


@router.put("/{document_id}", response_model=DocumentResponse)
async def update_document(
    document_id: int,
    data: DocumentUpdate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """Update a document. Leitung+ can edit."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    update_data = data.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(doc, field, value)

    log_action(db, current_user.id, "update", "document", doc.id, f"Dokument aktualisiert: {doc.titel}", request)
    db.commit()
    db.refresh(doc)
    return doc


@router.post("/{document_id}/upload", response_model=DocumentResponse)
async def upload_document_file(
    document_id: int,
    datei: UploadFile = File(...),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """Upload a file for a document. Leitung+ can upload."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    os.makedirs(DOCUMENT_UPLOAD_DIR, exist_ok=True)
    ext = _validate_upload(datei.filename)
    safe_filename = f"{uuid.uuid4().hex}{ext}"
    file_path = os.path.join(DOCUMENT_UPLOAD_DIR, safe_filename)

    content = await datei.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail=f"Datei zu groß. Maximum: {MAX_FILE_SIZE // (1024*1024)} MB")

    async with aiofiles.open(file_path, "wb") as f:
        await f.write(content)

    # Remove old file if exists
    if doc.datei_pfad and os.path.exists(doc.datei_pfad) and _safe_file_path(DOCUMENT_UPLOAD_DIR, doc.datei_pfad):
        os.remove(doc.datei_pfad)

    doc.datei_pfad = file_path
    db.commit()
    db.refresh(doc)
    return doc


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(
    document_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("admin")),
):
    """Dokument und Datei löschen. Nur Admin kann löschen."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    if doc.datei_pfad and os.path.exists(doc.datei_pfad) and _safe_file_path(DOCUMENT_UPLOAD_DIR, doc.datei_pfad):
        os.remove(doc.datei_pfad)

    # Delete associated aenderungsantrag
    db.query(DocumentAenderungsantrag).filter(DocumentAenderungsantrag.document_id == document_id).delete()

    doc_titel = doc.titel
    db.delete(doc)
    log_action(db, current_user.id, "delete", "document", document_id, f"Dokument gelöscht: {doc_titel}", request)
    db.commit()
    return None


# ---------------------------------------------------------------------------
# Aenderungsantrag CRUD
# ---------------------------------------------------------------------------

@router.get("/{document_id}/aenderungsantraege", response_model=List[DocumentAenderungsantragResponse])
async def list_aenderungsantraege(
    document_id: int,
    status_filter: Optional[str] = Query(None, alias="status"),
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """List aenderungsantraege for a document. Vorstand+ can view."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    query = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.document_id == document_id)
    )
    if status_filter:
        query = query.filter(DocumentAenderungsantrag.status == status_filter)
    return query.order_by(DocumentAenderungsantrag.created_at.desc()).all()


@router.post("/{document_id}/aenderungsantraege", response_model=DocumentAenderungsantragResponse, status_code=status.HTTP_201_CREATED)
async def create_aenderungsantrag(
    document_id: int,
    data: DocumentAenderungsantragCreate,
    send_emails: bool = Query(False, description="E-Mail-Benachrichtigung an konfigurierte Empfänger senden"),
    request: Request = None,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Create an amendment for a document. Vorstand+ can create. Optional: E-Mails an Benachrichtigungsempfänger senden."""
    doc = db.query(Document).filter(Document.id == document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")

    aenderungsantrag = DocumentAenderungsantrag(
        document_id=document_id,
        titel=getattr(data, "titel", None),
        antragsteller=data.antragsteller,
        antrag_text=data.antrag_text,
        alte_fassung=data.alte_fassung,
        neue_fassung=data.neue_fassung,
        begruendung=data.begruendung,
    )
    db.add(aenderungsantrag)
    db.flush()
    log_action(db, current_user.id, "create", "aenderungsantrag", aenderungsantrag.id, f"Änderungsantrag erstellt für Dokument {document_id}: {aenderungsantrag.antragsteller}", request)
    db.commit()
    db.refresh(aenderungsantrag)
    # Reload with stellen for response and optional email
    aenderungsantrag = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.id == aenderungsantrag.id)
        .first()
    )
    if send_emails:
        _send_aenderungsantrag_emails(db, aenderungsantrag, doc)
    return aenderungsantrag


@router.put("/aenderungsantraege/{aenderungsantrag_id}", response_model=DocumentAenderungsantragResponse)
async def update_aenderungsantrag(
    aenderungsantrag_id: int,
    data: DocumentAenderungsantragUpdate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """Update an aenderungsantrag. Leitung+ can edit (including status changes)."""
    aenderungsantrag = db.query(DocumentAenderungsantrag).filter(DocumentAenderungsantrag.id == aenderungsantrag_id).first()
    if not aenderungsantrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")

    update_data = data.model_dump(exclude_unset=True)

    if "status" in update_data:
        valid_statuses = ["eingereicht", "angenommen", "abgelehnt"]
        if update_data["status"] not in valid_statuses:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid status. Must be one of: {valid_statuses}",
            )

    for field, value in update_data.items():
        setattr(aenderungsantrag, field, value)

    log_action(db, current_user.id, "update", "aenderungsantrag", aenderungsantrag.id, f"Änderungsantrag aktualisiert: {aenderungsantrag.antragsteller}", request)
    db.commit()
    aenderungsantrag = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.id == aenderungsantrag_id)
        .first()
    )
    return aenderungsantrag


@router.post("/aenderungsantraege/{aenderungsantrag_id}/send-email")
async def send_aenderungsantrag_email(
    aenderungsantrag_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """E-Mail-Benachrichtigung für diesen Änderungsantrag an konfigurierte Empfänger senden. Leitung+."""
    antrag = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.id == aenderungsantrag_id)
        .first()
    )
    if not antrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag nicht gefunden")
    doc = db.query(Document).filter(Document.id == antrag.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Dokument nicht gefunden")
    if not settings.email_configured:
        raise HTTPException(status_code=503, detail="E-Mail ist nicht konfiguriert (SMTP).")
    if not settings.document_amendment_notify_emails_list:
        raise HTTPException(status_code=400, detail="Keine Empfänger konfiguriert (DOCUMENT_AMENDMENT_NOTIFY_EMAILS).")
    _send_aenderungsantrag_emails(db, antrag, doc)
    return {"detail": "E-Mails wurden versendet."}


# ---------------------------------------------------------------------------
# Stellen (einzelne Änderungen innerhalb eines Änderungsantrags)
# ---------------------------------------------------------------------------


@router.get("/aenderungsantraege/{aenderungsantrag_id}/stellen", response_model=List[DocumentAenderungResponse])
async def list_stellen(
    aenderungsantrag_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Liste aller Stellen eines Änderungsantrags. Vorstand+ kann lesen."""
    antrag = (
        db.query(DocumentAenderungsantrag)
        .filter(DocumentAenderungsantrag.id == aenderungsantrag_id)
        .first()
    )
    if not antrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")
    return sorted(antrag.stellen, key=lambda s: (s.position, s.id))


@router.post(
    "/aenderungsantraege/{aenderungsantrag_id}/stellen",
    response_model=DocumentAenderungResponse,
    status_code=status.HTTP_201_CREATED,
)
async def create_stelle(
    aenderungsantrag_id: int,
    data: DocumentAenderungCreate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Neue Stelle (Änderung) zu einem Änderungsantrag hinzufügen."""
    antrag = db.query(DocumentAenderungsantrag).filter(DocumentAenderungsantrag.id == aenderungsantrag_id).first()
    if not antrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")
    max_pos = max([s.position for s in antrag.stellen], default=-1)
    stelle = DocumentAenderung(
        aenderungsantrag_id=aenderungsantrag_id,
        position=data.position if data.position is not None else max_pos + 1,
        bezug=data.bezug,
        alte_fassung=data.alte_fassung,
        neue_fassung=data.neue_fassung,
        aenderungstext=data.aenderungstext,
    )
    db.add(stelle)
    db.flush()
    log_action(db, current_user.id, "create", "document_aenderung", stelle.id, f"Stelle hinzugefügt zu Antrag {aenderungsantrag_id}", request)
    db.commit()
    db.refresh(stelle)
    return stelle


@router.put("/aenderungsantraege/{aenderungsantrag_id}/stellen/{stelle_id}", response_model=DocumentAenderungResponse)
async def update_stelle(
    aenderungsantrag_id: int,
    stelle_id: int,
    data: DocumentAenderungUpdate,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("leitung")),
):
    """Stelle bearbeiten. Leitung+."""
    stelle = (
        db.query(DocumentAenderung)
        .filter(
            DocumentAenderung.id == stelle_id,
            DocumentAenderung.aenderungsantrag_id == aenderungsantrag_id,
        )
        .first()
    )
    if not stelle:
        raise HTTPException(status_code=404, detail="Stelle not found")
    for field, value in data.model_dump(exclude_unset=True).items():
        setattr(stelle, field, value)
    log_action(db, current_user.id, "update", "document_aenderung", stelle.id, f"Stelle {stelle_id} aktualisiert", request)
    db.commit()
    db.refresh(stelle)
    return stelle


@router.delete(
    "/aenderungsantraege/{aenderungsantrag_id}/stellen/{stelle_id}",
    status_code=status.HTTP_204_NO_CONTENT,
)
async def delete_stelle(
    aenderungsantrag_id: int,
    stelle_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("admin")),
):
    """Stelle löschen. Nur Admin."""
    stelle = (
        db.query(DocumentAenderung)
        .filter(
            DocumentAenderung.id == stelle_id,
            DocumentAenderung.aenderungsantrag_id == aenderungsantrag_id,
        )
        .first()
    )
    if not stelle:
        raise HTTPException(status_code=404, detail="Stelle not found")
    db.delete(stelle)
    log_action(db, current_user.id, "delete", "document_aenderung", stelle_id, f"Stelle gelöscht", request)
    db.commit()
    return None


# ---------------------------------------------------------------------------
# Export Änderungsantrag (DOCX + PDF: Änderungstext + Synopse)
# ---------------------------------------------------------------------------

@router.get("/aenderungsantraege/{aenderungsantrag_id}/export.docx")
async def export_aenderungsantrag_docx(
    aenderungsantrag_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Änderungsantrag als DOCX exportieren (Änderungstext + Synopse)."""
    antrag = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.id == aenderungsantrag_id)
        .first()
    )
    if not antrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")
    doc = db.query(Document).filter(Document.id == antrag.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    fd, path = tempfile.mkstemp(suffix=".docx")
    try:
        os.close(fd)
        _build_aenderungsantrag_docx(antrag, doc, path)
        with open(path, "rb") as f:
            docx_bytes = f.read()
        filename = f"aenderungsantrag_{aenderungsantrag_id}.docx"
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    finally:
        if os.path.exists(path):
            try:
                os.unlink(path)
            except OSError:
                pass


@router.get("/aenderungsantraege/{aenderungsantrag_id}/export.pdf")
async def export_aenderungsantrag_pdf(
    aenderungsantrag_id: int,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("vorstand")),
):
    """Änderungsantrag als PDF exportieren (DOCX wird mit LibreOffice konvertiert)."""
    antrag = (
        db.query(DocumentAenderungsantrag)
        .options(joinedload(DocumentAenderungsantrag.stellen))
        .filter(DocumentAenderungsantrag.id == aenderungsantrag_id)
        .first()
    )
    if not antrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")
    doc = db.query(Document).filter(Document.id == antrag.document_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    fd, path = tempfile.mkstemp(suffix=".docx")
    try:
        os.close(fd)
        _build_aenderungsantrag_docx(antrag, doc, path)
        pdf_path = await asyncio.to_thread(docx_to_pdf, path)
        if not pdf_path or not os.path.isfile(pdf_path):
            raise HTTPException(status_code=500, detail="PDF-Konvertierung fehlgeschlagen (LibreOffice erforderlich).")
        filename = f"aenderungsantrag_{aenderungsantrag_id}.pdf"
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )
    finally:
        for p in (path, path.replace(".docx", ".pdf")):
            if os.path.exists(p):
                try:
                    os.unlink(p)
                except OSError:
                    pass


@router.delete("/aenderungsantraege/{aenderungsantrag_id}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_aenderungsantrag(
    aenderungsantrag_id: int,
    request: Request,
    db: Session = Depends(get_db),
    current_user: User = Depends(require_role("admin")),
):
    """Änderungsantrag löschen. Nur Admin kann löschen."""
    aenderungsantrag = db.query(DocumentAenderungsantrag).filter(DocumentAenderungsantrag.id == aenderungsantrag_id).first()
    if not aenderungsantrag:
        raise HTTPException(status_code=404, detail="Aenderungsantrag not found")

    antragsteller = aenderungsantrag.antragsteller
    db.delete(aenderungsantrag)
    log_action(db, current_user.id, "delete", "aenderungsantrag", aenderungsantrag_id, f"Änderungsantrag gelöscht: {antragsteller}", request)
    db.commit()
    return None
